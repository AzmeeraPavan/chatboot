"""
Improved RAG Pipeline
-----------------------------------------
✓ Excel row-level indexing
✓ Hybrid Search (BM25 + Embeddings)
✓ Exact-match boosting
✓ Correct LangChain imports
✓ Works with txt/csv/xlsx/pdf/docx/json
"""

import os
import sys
import json
from pathlib import Path
from typing import List
import pandas as pd
from pypdf import PdfReader
import docx

# LangChain (updated imports)
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.retrievers import BM25Retriever

# HuggingFace LLM
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline

# llama.cpp
try:
    from llama_cpp import Llama
    LLAMA_AVAILABLE = True
except Exception:
    LLAMA_AVAILABLE = False


# ---------------- CONFIG ----------------
DATA_DIR = r"C:\Users\PavanAzmeeraSIDGloba\rowdata"
VECTOR_DIR = r"C:\vectorData\langchain_faiss"
EMBED_MODEL = "sentence-transformers/all-mpnet-base-v2"
CHUNK_SIZE = 400
CHUNK_OVERLAP = 40
TOP_K = 1
HF_MODEL = "google/gemma-2b-it"
LLAMA_MODEL_PATH = "models/llama-3-8b.gguf"
# ----------------------------------------


# ====================================================
#  Excel Loader  (each row becomes one Document)
# ====================================================
def load_excel_as_rows(path: Path) -> List[Document]:
    df = pd.read_excel(path, dtype=str).fillna("")
    docs = []

    for _, row in df.iterrows():
        row_text = " | ".join(f"{col}: {row[col]}" for col in df.columns)
        docs.append(Document(
            page_content=row_text,
            metadata={"source": str(path), "row_id": int(_)}
        ))

    return docs


# ====================================================
# Standard File Readers
# ====================================================
def read_txt(p: Path): return p.read_text(encoding="utf-8", errors="ignore")

def read_csv(p: Path):
    df = pd.read_csv(p, dtype=str).fillna("")
    return df.to_string(index=False)

def read_pdf(p: Path):
    reader = PdfReader(str(p))
    pages = [pg.extract_text() or "" for pg in reader.pages]
    return "\n".join(pages)

def read_docx(p: Path):
    d = docx.Document(str(p))
    return "\n".join([p.text for p in d.paragraphs])

def read_json(p: Path):
    return json.dumps(json.load(open(p, encoding="utf-8")), indent=2, ensure_ascii=False)

READERS = {
    ".txt": read_txt,
    ".csv": read_csv,
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".json": read_json
}


# ====================================================
# Build Vector Store (with hybrid search)
# ====================================================
def build_vectorstore(data_dir=DATA_DIR, vector_dir=VECTOR_DIR):
    files = list(Path(data_dir).rglob("*"))
    print(f"[INFO] Found {len(files)} files")

    docs = []

    for p in files:
        if not p.is_file():
            continue

        ext = p.suffix.lower()

        # Excel is handled differently (row-level)
        if ext in [".xlsx", ".xls"]:
            docs.extend(load_excel_as_rows(p))
            continue

        # Other files
        if ext in READERS:
            txt = READERS[ext](p)
            if txt.strip():
                docs.append(Document(page_content=txt, metadata={"source": str(p)}))

    print(f"[INFO] Loaded {len(docs)} documents")

    # Text splitter for large non-Excel files
    splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
    chunks = splitter.split_documents(docs)
    print(f"[INFO] Split into {len(chunks)} chunks")

    embedder = HuggingFaceEmbeddings(model_name=EMBED_MODEL)

    db = FAISS.from_documents(chunks, embedder)
    os.makedirs(vector_dir, exist_ok=True)
    db.save_local(vector_dir)

    print(f"[SUCCESS] Saved FAISS vectorstore to: {vector_dir}")


# ====================================================
# Hybrid Retrieval (BM25 + Embeddings + Exact Match)
# ====================================================
def load_vectorstore(vector_dir=VECTOR_DIR, embed_model=EMBED_MODEL):
    embedder = HuggingFaceEmbeddings(model_name=embed_model)
    db = FAISS.load_local(vector_dir, embedder, allow_dangerous_deserialization=True)
    return db


def retrieve(query: str, top_k=TOP_K):
    db = load_vectorstore()

    # Embedding-based retriever
    emb_docs = db.similarity_search(query, k=top_k)

    # BM25 keyword retriever (text only)
    all_docs = db.docstore._dict.values()
    bm25 = BM25Retriever.from_documents(list(all_docs))
    bm_docs = bm25._get_relevant_documents(query, run_manager=None)

    # Exact-match boost
    exact = [
        d for d in all_docs
        if query.lower() in d.page_content.lower()
    ]

    # Merge & de-duplicate
    merged = exact + emb_docs + bm_docs
    final = []
    seen = set()

    for d in merged:
        key = (d.metadata.get("source"), d.page_content[:30])
        if key not in seen:
            final.append(d)
            seen.add(key)

        if len(final) >= top_k:
            break

    return final


# ====================================================
#  GENERATION LOGIC
# ====================================================
def generate_with_llama(context, query):
    if not LLAMA_AVAILABLE:
        return "[ERROR: llama.cpp not installed]"
    llm = Llama(model_path=LLAMA_MODEL_PATH, n_ctx=4096)
    prompt = f"Context:\n{context}\nQuestion: {query}\nAnswer:"
    resp = llm(prompt, max_tokens=300)
    return resp["choices"][0]["text"]


def generate_with_hf(context, query):
    tok = AutoTokenizer.from_pretrained(HF_MODEL)
    model = AutoModelForCausalLM.from_pretrained(HF_MODEL)
    gen = pipeline("text-generation", model=model, tokenizer=tok)

    prompt = f"Use ONLY this context:\n{context}\nQ: {query}\nA:"
    out = gen(prompt, max_new_tokens=200, do_sample=False)
    return out[0]["generated_text"]


# ====================================================
# CLI Commands
# ====================================================
def cmd_build():
    build_vectorstore()


def cmd_retrieve(q):
    docs = retrieve(q)
    print(f"\n=== Retrieved {len(docs)} Chunks ===\n")
    for i, d in enumerate(docs, 1):
        print(f"--- CHUNK {i} ({d.metadata.get('source')}) ---\n{d.page_content}\n")


def cmd_answer(q, engine="llama"):
    docs = retrieve(q)
    if not docs:
        print("[No results]")
        return

    context = "\n\n".join(d.page_content for d in docs)

    if engine == "hf":
        ans = generate_with_hf(context, q)
    else:
        ans = generate_with_llama(context, q)

    print("\n===== ANSWER =====\n")
    print(ans)
    print("\n===================\n")


# ====================================================
# MAIN
# ====================================================
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python rag_pipeline.py [build|retrieve|answer] <query>")
        sys.exit()

    cmd = sys.argv[1].lower()

    if cmd == "build":
        cmd_build()

    elif cmd == "retrieve":
        q = " ".join(sys.argv[2:])
        cmd_retrieve(q)

    elif cmd == "answer":
        q = " ".join(arg for arg in sys.argv[2:] if not arg.startswith("--"))
        engine = "hf" if "--hf" in sys.argv else "llama"
        cmd_answer(q, engine)

    else:
        print("Unknown command.")
