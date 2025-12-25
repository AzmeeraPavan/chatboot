import json
import pandas as pd
from pathlib import Path
from typing import List
import docx
from pypdf import PdfReader
from langchain_core.documents import Document

def load_excel_as_rows(path: Path) -> List[Document]:
    df = pd.read_excel(path, dtype=str).fillna("")
    docs = []

    for idx, row in df.iterrows():
        text = " | ".join(f"{col}: {row[col]}" for col in df.columns)
        docs.append(Document(
            page_content=text,
            metadata={"source": str(path), "row_id": int(idx)}
        ))
    return docs

def read_txt(p): return p.read_text(encoding="utf-8", errors="ignore")

def read_csv(p):
    df = pd.read_csv(p, dtype=str).fillna("")
    return df.to_string(index=False)

def read_pdf(p):
    reader = PdfReader(str(p))
    pages = [pg.extract_text() or "" for pg in reader.pages]
    return "\n".join(pages)

def read_docx(p):
    d = docx.Document(str(p))
    return "\n".join([p.text for p in d.paragraphs])

def read_json(p):
    return json.dumps(json.load(open(p, encoding="utf-8")), indent=2, ensure_ascii=False)

READERS = {
    ".txt": read_txt,
    ".csv": read_csv,
    ".pdf": read_pdf,
    ".docx": read_docx,
    ".json": read_json
}

def load_documents(data_dir: str) -> List[Document]:
    files = list(Path(data_dir).rglob("*"))
    print(f"[INFO] Found {len(files)} files in {data_dir}")

    docs = []

    for p in files:
        if not p.is_file():
            continue

        ext = p.suffix.lower()

        if ext in [".xlsx", ".xls"]:
            docs.extend(load_excel_as_rows(p))
            continue

        if ext in READERS:
            try:
                text = READERS[ext](p)
                if text.strip():
                    docs.append(Document(page_content=text, metadata={"source": str(p)}))
            except Exception as e:
                print(f"[WARN] Could not read {p}: {e}")

    return docs
